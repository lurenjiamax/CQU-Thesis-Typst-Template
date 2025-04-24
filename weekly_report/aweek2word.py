import re
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# 解析 typst 代码并提取内容
def extract_typst_content(typst_code):
    # 移除以 // 开头的注释行
    typst_code = re.sub(r'//.*', '', typst_code)
    typst_code = str(typst_code).replace("*", "")

    # 提取周次
    week_match = re.search(r'#let week = (\d+)', typst_code)
    week = int(week_match.group(1)) if week_match else 1  # 默认值为1
    # print(typst_code)
    # 提取各个部分内容
    work_plan = re.search(r'== 一、本周计划及执行情况\n(.*?)\n==', typst_code, re.S)
    work = re.search(r'== 二、实际工作\n(.*?)\n==', typst_code, re.S)
    problem = re.search(r'== 三、存在问题\n(.*?)\n==', typst_code, re.S)
    next_week_plan = re.search(r'== 四、下周计划\n(.*)$', typst_code, re.S)
    # print(next_week_plan.group(1))
    return {
        "week": week,
        "work_plan": work_plan.group(1).strip() if work_plan else "",
        "work": work.group(1).strip() if work else "",
        "problem": problem.group(1).strip() if problem else "",
        "next_week_plan": next_week_plan.group(1).strip() if next_week_plan else "",
    }

# 示例 Typst 代码
def load_code(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        file = file.read()
    return extract_typst_content(file)

# 加载 Word 模板
def load_template(template_path):
    document = Document(template_path)
    style = document.styles
    obj_charstyle = style.add_style('SimSun_body_style', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'SimSun'
    
    obj_charstyle = style.add_style('SimSun_bold_style', WD_STYLE_TYPE.PARAGRAPH)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'SimSun'
    obj_font.bold = True
    return document

# 替换模板中的占位符
def replace_placeholder(doc, placeholder, replacement):
    for para in doc.paragraphs:
        if placeholder in para.text:
            if placeholder == "{{week}}":
                t = f"工作时间段：第{replacement}周"
                para.clear()
                run = para.add_run(t)
                para.style = doc.styles['SimSun_bold_style']
                run.font.name = 'SimSun'
                run.font.size = Pt(12)
                run.bold = True
                # para.text = para.text.replace(placeholder, str(replacement))
                # para.style = doc.styles['SimSun_bold_style']
            else:
                # para.text = para.text.replace(placeholder, str(replacement))
                # para.style = doc.styles['SimSun_body_style']
                para.clear()
                run = para.add_run(replacement)
                para.style = doc.styles['SimSun_body_style']
                run.font.name = 'SimSun'
                run.font.size = Pt(12)

def replace_warper(doc, extracted_content):
    replace_placeholder(doc, "{{week}}", extracted_content["week"])
    replace_placeholder(doc, "{{work_plan}}", extracted_content["work_plan"])
    replace_placeholder(doc, "{{work}}", extracted_content["work"])
    replace_placeholder(doc, "{{problem}}", extracted_content["problem"])
    replace_placeholder(doc, "{{next_week_plan}}", extracted_content["next_week_plan"])
    doc.save("周记第{}周.docx".format(extracted_content["week"]))


def all(file_path, template_path):
    extracted_content = load_code(file_path)
    doc = load_template(template_path)
    replace_warper(doc, extracted_content)
    


for week in range(8, 10):
    all(f"week{week}.typ", "周记模板.docx")